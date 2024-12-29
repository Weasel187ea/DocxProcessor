import csv
import re
import docx
import os

# from translate import Translator
# translator= Translator(from_lang="german",to_lang="english")
# # translation = translator.translate("Guten Morgen")


def write_into_file(data): # write into file in csv format
    with open("finalData.csv", "a", newline="") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writerows([data])


def iter_headings(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading'):
            yield paragraph

def ignoreLine(text):
    if not text: #if line is empty, ignore
        return True
    elif ("Load-Date:" in text) or ("End of Document" in  text): #if includes either, we do not add to body text
        return True
    elif ("Body" in text) or ("PDF" in text):
        return True
    return False

def getText(filename):
    doc = docx.Document(filename)
    fullText = [] #all text will be joined at the end
    mainHeader = "" #header
    body = False #we have reached the body paragraph, collect everything
    headerData = []
    headerDataDone = False
    importWords = set() #small feature, includes words that are bolded in the body of text
    Data = {} # --> [header, length, section, text]


    for para in doc.paragraphs:

        if mainHeader == "": #this section extracts header
            for header in iter_headings(doc.paragraphs):
                Data["title"] = header.text

        for run in para.runs:
            if run.bold:
                # if "Section" in run.text:
                if re.findall(r"Section:", run.text):
                    Data["section"] = para.text.replace(u'\xa0', u' ').replace("\n"," ").replace("Section:", "")
                elif re.findall(r"Length:", run.text):
                    Data["length"] = para.text.replace(u'\xa0', u' ').replace("\n"," ").replace("Length:", "")
                elif re.findall(r"Byline:", run.text):
                    Data["byline"] = para.text.replace(u'\xa0', u' ').replace("\n"," ").replace("Byline:", "")
                elif re.findall(r"Body", run.text):
                    body = True
                elif run.font.underline:
                    importWords.add(run.text)

                elif (":" not in run.text) and (run.text not in ["Load-Date:", "End of Document"]) and (run.text[-1].isalpha()):
                        # importWords.append(run.text) # keep track of important words
                    importWords.add(run.text)
                    

                headerDataDone = True # Once we reach bold words we can assume header info has been processed
                if len(headerData) >= 1:
                    if ":" in headerData[1]:
                        headerData[1] = " ".join(headerData[1].split(":")[1:])

                    Data["publisher"] = headerData[1]
                if len(headerData) >=2:
                    Data["date"] = headerData[2]
                

        if (not headerDataDone) and (para.text != ""): #data in header is not part of body paragraph, stored differently
            headerData.append(para.text)

        if ignoreLine(para.text):
            continue

        elif body == True: #we are in the body paragraph
                if (para.text != mainHeader) and (para.text != ""): #check that text is not header text nore empty
                    lineToAdd = para.text.replace("\n", " ") #remove newlines
                    fullText.append(lineToAdd) #add to fullTest list, will be joined later

        
    Data["body"] = (" ".join(fullText)) 
    Data["bold"] = ", ".join(importWords)
    #########################################  section that writes to csv file
    write_into_file(Data)
    #########################################
    return Data



if __name__ == "__main__":

    with open("finalData.csv", "w") as csvfile:
        fieldnames = ["title", "byline", "date", "section", "publisher", "length", "body", "bold"]
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader() #initializes header
        


    onlyfiles = os.listdir() #extracts all files from directory that are docx formatted

    for file in os.listdir("Documents"):
        f = os.path.join("Documents", file) #path to file
        if f.endswith("docx"):
            getText(f)